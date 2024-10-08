import base64
import io
import re
import json
import boto3
from PIL import Image
from botocore.exceptions import ClientError
from enum import Enum
from io import BytesIO
import sagemaker
from typing import Any, List
from langchain_core.pydantic_v1 import BaseModel, Field
import os
from sagemaker.async_inference.waiter_config import WaiterConfig
import time
from sagemaker.predictor_async import AsyncPredictor
from sagemaker.serializers import JSONSerializer
from sagemaker.deserializers import JSONDeserializer
from sagemaker.predictor import Predictor
from docx import Document
from docx.shared import Inches
from docx2pdf import convert

class StyleEnum(Enum):
    Photographic = "photographic"
    Tile_texture = "tile-texture"
    Digital_art = "digital-art"
    Origami = "origami"
    Modeling_compound = "modeling-compound"
    Anime = "anime"
    Cinematic = "cinematic"
    Model_3D = "3d-model"
    Comicbook = "comic-book"
    Enhance = "enhance"
    
class ImageError(Exception):
    "Custom exception for errors returned by SDXL"
    def __init__(self, message):
        self.message = message


class ImageGenerator(BaseModel):
    """
        invoke SDXL model in a Amaozn Bedrock to generate identity images

    """
    model_id: str = Field(default="amazon.titan-image-generator-v2:0")
    cfg_scale: int = Field( default=7)
    steps:int = Field( default=50)
    samples:int = Field( default=1)
    
    def _generate(self,model_id, body):
        """
        Generate an image using SDXL 1.0 on demand.
        Args:
            model_id (str): The model ID to use.
            body (str) : The request body to use.
        Returns:
            image_bytes (bytes): The image generated by the model.
        """

        # logger.info("Generating image with SDXL model %s", model_id)

        session = boto3.Session()
        #get bedrock service 
        bedrock = session.client(service_name='bedrock-runtime')
    
        accept = "application/json"
        content_type = "application/json"

        response = bedrock.invoke_model(
            body=body, modelId=model_id, accept=accept, contentType=content_type
        )
        response_body = json.loads(response.get("body").read())
        
        if model_id.startswith("stability"):
            base64_image = response_body.get("artifacts")[0].get("base64")
            base64_bytes = base64_image.encode('ascii')
            image_bytes = base64.b64decode(base64_bytes)

            finish_reason = response_body.get("artifacts")[0].get("finishReason")

            if finish_reason == 'ERROR' or finish_reason == 'CONTENT_FILTERED':
                raise ImageError(f"Image generation error. Error code is {finish_reason}")
        else:
            base64_image = response_body.get("images")[0]
            base64_bytes = base64_image.encode('ascii')
            image_bytes = base64.b64decode(base64_bytes)

            finish_reason = response_body.get("error")

            if finish_reason is not None:
                raise ImageError(f"Image generation error. Error is {finish_reason}")

        print(f"Successfully generated image with model {model_id}")

        return image_bytes

    def generate_image( self,prompt,seed=0,style_preset=StyleEnum.Photographic.value):
        if self.model_id.startswith('stability'):
            body=json.dumps({
                "text_prompts": [
                {
                "text": prompt
                }
            ],
            "cfg_scale": self.cfg_scale,
            "seed": seed,
            "steps": self.steps,
              "height": 768,
            "width": 768,
            "samples" : self.samples,
            "style_preset" : style_preset
            })
        elif self.model_id.startswith('amazon'):
            body = json.dumps({
            "taskType": "TEXT_IMAGE",
            "textToImageParams": {
                "text": prompt
            },
            "imageGenerationConfig": {
                "numberOfImages": 1,
                "height": 768,
                "width": 768,
                "cfgScale": self.cfg_scale,
                "seed": seed
            }
            })
        print(body)
        image= None
        try:
            image_bytes=self._generate(model_id = self.model_id,
                                    body = body)
            image = Image.open(io.BytesIO(image_bytes))

        except ClientError as err:
            message=err.response["Error"]["Message"]
            # logger.error("A client error occurred: %s", message)
            print("A client error occured: " +format(message))
        except ImageError as err:
            print(err)
        except Exception as err:
            print(err)
        finally:
            return image
        



def save_image_file(image, filename, folder):
    if not os.path.exists(folder):
        os.makedirs(folder)
    image.save(os.path.join(folder, filename))
    print(f"image saved in {os.path.join(folder, filename)}")


def save_all_images(images,folder='./images'):
    img_fnames = []
    for i,img in enumerate(images):
        fname = str(i)+'.png'
        if img:
            save_image_file(img,fname,folder)
            img_fnames.append(os.path.join(folder, fname))
        else:
            img_fnames.append(None)
    return img_fnames

def save_all_images_names(images:list,file_names:list,folder:str ='./images'):
    img_fnames = []
    for img,fname in zip(images,file_names):
        fname =fname +'.png'
        if img:
            save_image_file(img,fname,folder)
            img_fnames.append(os.path.join(folder, fname))
        else:
            img_fnames.append(None)
    return img_fnames


def base64_to_image(base64_string):
    image_bytes = base64.b64decode(base64_string)
    image_buffer = BytesIO(image_bytes)
    image = Image.open(image_buffer)
    return image


def Image2base64(img_path):
    image = Image.open(img_path)
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    image_data = buffer.getvalue()
    base64_encoded_string = base64.b64encode(image_data).decode('utf-8')
    return base64_encoded_string

def get_bucket_and_key(s3uri):
    pos = s3uri.find("/", 5)
    bucket = s3uri[5:pos]
    key = s3uri[pos + 1 :]
    return bucket, key

class StoryDiffusionGenerator():
    """
        invoke storydiffusion model hosted in a SageMaker endpoint to generate consistant images
    """
    
    def __init__(self,endpoint_name):
        endpoint_name = endpoint_name
        # boto_session= boto3.Session(profile_name=profile)
        boto_session= boto3.Session()
        self.s3_resource = boto_session.resource("s3")
        sagemaker_session = sagemaker.Session(boto_session = boto_session)
        bucket  = sagemaker_session.default_bucket()
        output_path  = "s3://{0}/{1}/asyncinvoke/out/".format(bucket, "story-diffusion")
        input_path :str = "s3://{0}/{1}/asyncinvoke/in/".format(bucket, "story-diffusion")
        
        predictor_ = Predictor(
            endpoint_name=endpoint_name,
            sagemaker_session=sagemaker_session,
            model_data_input_path=input_path,
            model_data_output_path=output_path,
        )
        predictor_.serializer = JSONSerializer()
        predictor_.deserializer = JSONDeserializer()
        self.config = WaiterConfig(
            max_attempts=100, delay=10  #  number of attempts  #  time in seconds to wait between attempts
        )
        self.predictor_async = AsyncPredictor(
                predictor_,
                name='story-diffusion'
        )
    
    def generate_real_identity_images(self,prompt:str, general_prompt:str = '', height:int = 768, width :int = 768):
        images = self.generate_images(general_prompt = general_prompt,
                                                            style="Photographic",
                                                            comic_type = "Classic Comic Style",
                                                            prompt_array=prompt,
                                                            id_length= 0,
                                                            sd_type = "Unstable",
                                                            ref_imgs=[],height=height,width=width) 
        for img in images:
            if img.size[0] < 1024:
                return img
        return None
        
    def generate_images(self,general_prompt:str,prompt_array:str,id_length:int=2, ref_imgs: List[Any]= [],comic_type:str='Classic Comic Style', style:str = 'Japanese Anime',sd_type:str="Unstable", height:int = 768, width :int = 768) -> list:
        data = { "general_prompt": general_prompt,
                        "prompt_array" : prompt_array,
                        "style" : style,
                        "G_height" : height,
                        "G_width" : width,
                        "comic_type" : comic_type,
                       "files":ref_imgs,
                        "id_length_":id_length,
                        "sd_type":sd_type,
                }
        if not ref_imgs:
            del data['files']
        # print(data)
        prediction = self.predictor_async.predict_async(data)
        print(f"Response output path: {prediction.output_path}")
        start = time.time()
        prediction.get_result(self.config)
        print(f"Time taken: {time.time() - start}s")
        
        output_bucket, output_key = get_bucket_and_key(prediction.output_path)
        output_obj = self.s3_resource.Object(output_bucket, output_key)
        body = output_obj.get()["Body"].read().decode("utf-8")
        
        respobj = json.loads(body)
        images = []
        for img in respobj['images_base64']:
            images.append(base64_to_image(img))
            
        return images
    

# each story line will send to storydiffusion model to create a comic, count the characters in each line and add crespondant ref images
def count_character_names(character_names,line):
    name_counter = {}
    for name in character_names:
        if f"[{name}]" in line:
            if name in name_counter:
                name_counter[name] += 1
            else:
                name_counter[name] = 1
    return name_counter if name_counter else {'[NC]':1}


#https://github.com/HVision-NKU/StoryDiffusion/blob/main/utils/gradio_utils.py
# convert character list to dict
def character_to_dict(general_prompt):
    character_dict = {}    
    generate_prompt_arr = general_prompt.splitlines()
    character_index_dict = {}
    invert_character_index_dict = {}
    character_list = []
    for ind,string in enumerate(generate_prompt_arr):
        # 分割字符串寻找key和value
        start = string.find('[')
        end = string.find(']')
        if start != -1 and end != -1:
            key = string[start:end+1]
            value = string[end+1:]
            if "#" in value:
                value =  value.rpartition('#')[0] 
            if key in character_dict:
                raise Exception("duplicate character descirption: " + key)
            character_dict[key] = value
            character_list.append(key)

        
    return character_dict 

def calc_id_length_prompt(general_prompt,prompts):
    character_dict = character_to_dict(general_prompt)
    replace_prompts = []
    character_index_dict = {}
    invert_character_index_dict = {}
    for ind,prompt in enumerate(prompts):
        for key in character_dict.keys():
            if key in prompt:
                if key not in character_index_dict:
                    character_index_dict[key] = []
                character_index_dict[key].append(ind)
                if ind not in invert_character_index_dict:
                    invert_character_index_dict[ind] = []
                invert_character_index_dict[ind].append(key)
        cur_prompt = prompt
        if ind in invert_character_index_dict:
            for key in invert_character_index_dict[ind]:
                cur_prompt = cur_prompt.replace(key,character_dict[key])
        replace_prompts.append(cur_prompt)
    # print(replace_prompts)
    ref_index_dict = {}
    ref_totals = []
    # print(character_index_dict)
    id_length = 999
    for character_key in character_index_dict.keys():
        if character_key not in character_index_dict:
            raise Exception("{} not have prompt description, please remove it".format(character_key))
        index_list = character_index_dict[character_key]
        # print(invert_character_index_dict)
        index_list = [index for index in index_list if len(invert_character_index_dict[index]) == 1]
        # print(f'{character_key}:',index_list)
        id_length = len(index_list) if len(index_list) < id_length else id_length
        # if len(index_list) < id_length:
        #     raise Exception(f"{character_key} not have enough prompt description, need no less than {id_length}, but you give {len(index_list)}")
        ref_index_dict[character_key] = index_list[:id_length]
        ref_totals = ref_totals + index_list[:id_length]
    return id_length
    
def generate_img_dicts(characters):
    character_names = [characters.main_character.name]
    name_figure_map = {characters.main_character.name:characters.main_character.figure}
    
    # add supporting characters
    for ch in characters.supporting_character:
        character_names += [ch.name]
        name_figure_map[ch.name] = ch.figure
        

    # print('name_figure_map:',name_figure_map)
    imgs = {}
    for key in list(name_figure_map.keys()):
        if key not in imgs :
            imgs[key] = Image2base64(f'./images/{key}.png')
    return imgs

    
def prepare_storyd_prompts(story_lines,characters,img_dicts):
    """
        prepare prompts for story diffusion model. 
        Use character's portrait as reference images to keep the consistance
    """
    character_names = [characters.main_character.name]
    name_figure_map = {characters.main_character.name:characters.main_character.figure}
    
    # add supporting characters
    for ch in characters.supporting_character:
        character_names += [ch.name]
        name_figure_map[ch.name] = ch.figure
        
    # count character names in each line
    name_counters = []
    for line in story_lines:
        name_counters.append(count_character_names(character_names,line))
    # print('name_counters:',name_counters)
    # generate prompt for each line
    args = []
    for name_counter,line in zip(name_counters,story_lines):
        # id_length = len(list(name_counter.keys()))
        ref_imgs = []
        figures = []
        for key in list(name_counter.keys()):
            if key != '[NC]':
                ref_imgs.append(img_dicts[key])
                figures.append(f"[{key}] {name_figure_map[key]} img")
            else:
                figures.append(f"[NC]")
                
        prompt_array = line.split("\n")

        # the model cannot generate one image with more than 2 character identities
        prompt_array_new = []
        pattern = r"\[(.*?)\]"
        for text in prompt_array:
            new_prompt = text
            match = re.findall(pattern, text)
            # check if the identity is in the first two of character dict
            if match and match[0] != 'NC' and match[0] not in list(name_counter.keys())[:2]:
                new_prompt = new_prompt.replace(f"[{match[0]}]",name_figure_map[key],1) 
                print(match[0], name_figure_map.keys(), name_counter.keys(),text,new_prompt)

            match = re.findall(pattern, new_prompt)
            if match and len(match) > 1:
                for k in match[1:]:
                    if k != 'NC' and k in list(name_counter.keys()): 
                        new_prompt = new_prompt.replace(f"[{k}]",name_figure_map[k]) 

            new_prompt = new_prompt if new_prompt.startswith('[NC]') else new_prompt

            match = re.findall(pattern, new_prompt)
            if not match:# if there is no bracket in the prompt line, it has to be add [NC]
                new_prompt = '[NC]' + new_prompt 
            
            prompt_array_new.append(new_prompt +'#' + text )# the text after # becomes caption 

        # calc id length
        general_prompt = '\n'.join(figures[:2]) #can only accept the first 2 figures currently, need to update in future              

        id_length = 2
        # add extral prompt for identity in case have enough prompt description for identity
        prompt_array_new = [f.replace(' img','') for f in figures[:id_length]] + prompt_array_new

        #re calc again after the prompt changed
        id_length = calc_id_length_prompt(general_prompt, prompt_array_new)
        id_length = 2 if id_length > 2 else id_length
        # print(prompt_array_new)    
        
        # now the model can only support max 2 ref images in general prompt
        # args.append({'prompt_array':prompt_array_new,'id_length':id_length,'ref_imgs':ref_imgs[:2],'general_prompt':general_prompt})
        yield({'prompt_array':prompt_array_new,'id_length':id_length,'ref_imgs':ref_imgs[:2],'general_prompt':general_prompt})
    # return args
    
    


def save_image( image, folder='./images') -> str:
    """Save the image to a temporary file and return the file path."""
    image_path = f"temp_{hash(image.tobytes())}.png"
    filename = os.path.join(folder,image_path)
    image.save(filename)
    return filename

def save_as_docx(characters,story, fname, suffix='_refined'):
    document = Document()
    document.add_heading(story.story_title, 0)
    document.add_heading('Characters introduction', level=1)
    character = characters.main_character
    document.add_heading(f"{character.name}", level=2)
    document.add_paragraph(f"Role: {character.role}\nBackground: {character.background}")
    document.add_picture(story.identity_images[0], width=Inches(4))
    for character,id_img in zip(characters.supporting_character,story.identity_images):
        document.add_heading(f"{character.name}", level=2)
        document.add_paragraph(f"Role: {character.role}\nBackground: {character.background}")
        document.add_picture(id_img, width=Inches(4))

    img_idx = 0
    for chapter in story.chapters:
        document.add_heading(chapter.chapter_title, level=1)
        para = chapter
        print(para)
        images = story.images[img_idx]
        img_idx += 1
        document.add_paragraph(para.content)

        # Add the image
        if images:
            [document.add_picture(image, width=Inches(6)) for image in images] 
        else:
            document.add_picture('placeholder.png', width=Inches(6))

    # Save the document
    document.save(fname)
    print(f'docx file saved as: {fname}')
    
    #Convert the Word document to PDF
    # pdf_path = f"{fname}.pdf"
    # convert(fname, pdf_path)
    # print(f'pdf file saved as: {pdf_path}')



def save_as_docx_old(characters,story, fname,suffix=''):
    document = Document()

    document.add_heading(story.story_title, 0)
    document.add_heading('Characters introduction', level=1)
    character = characters.main_character
    document.add_heading(f"{character.name}", level=2)
    document.add_paragraph(f"Role: {character.role}\nBackground: {character.background}")
    document.add_picture(f'./images/{character.name}{suffix}.png', width=Inches(4))
    for character in characters.supporting_character:
        document.add_heading(f"{character.name}", level=2)
        document.add_paragraph(f"Role: {character.role}\nBackground: {character.background}")
        document.add_picture(f'./images/{character.name}{suffix}.png', width=Inches(4))

    img_idx = 0
    for chapter in story.chapters:
        document.add_heading(chapter.chapter_title, level=1)
        # Add the image
        image = story.images[img_idx]
        img_idx += 1
        if image:
            document.add_picture(save_image(image), width=Inches(6))
        else:
            document.add_picture('placeholder.png', width=Inches(6))
        # Add paragraphs
        for para in chapter.paragraphs:
            print(para)
            document.add_paragraph(para.content)

    # Save the document
    document.save(fname)
    print(f'docx file saved as: {fname}')
    
    # Convert the Word document to PDF
    pdf_path = f"{fname}.pdf"
    try:
        convert(fname, pdf_path)
        print(f'pdf file saved as: {pdf_path}')
    except Exception as e:
        print(e)
        print('pdf file not saved')